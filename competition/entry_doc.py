"""Word (.docx) rendering of a player's game-week entry.

Mirrors the paper "Charity Tipster" form (see exampleform.pdf) so the online
half of the game can print/email an identical sheet: the four sections, the
player's predictions filled in, and the scoring boxes left blank for marking.

`build_entry_docx()` returns the .docx as bytes (built in memory — nothing hits
disk, which matters on serverless). `entry_completeness()` is the gate the views
and template use to only offer download/email once the form is actually filled.
"""

from __future__ import annotations

import io

from django.utils import timezone

from .models import Entry

# Scoring legends, lifted verbatim from the paper form.
_SECTION1_LEGEND = (
    "3pts Home Win, 4pts Draw, 5pts Away Win plus [5pts Bonus for Correct Result "
    "totalling 0–4 goals or 7pts for Correct Result totalling 5 or more goals]."
)
_SECTION2_LEGEND = (
    "(5 points exactly correct: 3 pts for +/- 1 Goal: 2 pts for +/- 2 Goals: "
    "1 pt for +/- 3 Goal)"
)
_SECTION4_POINTS = {1: "4 Points", 2: "3 Points", 3: "2 Points", 4: "1 Point"}


def entry_completeness(participant, game_week):
    """Return (entry, missing) where `missing` is a list of human labels.

    The form counts as complete when every match has both scores, the total-goals
    box is filled, and all four scorers are named. True/False always has a value
    (it defaults to True), so it never blocks. `entry` is None when nothing has
    been saved yet.
    """
    entry = (
        Entry.objects.filter(participant=participant, game_week=game_week)
        .prefetch_related("match_predictions", "scorer_picks")
        .first()
    )
    if entry is None:
        return None, ["Save your predictions first"]

    missing = []
    preds = {mp.fixture_id: mp for mp in entry.match_predictions.all()}
    unscored = 0
    for fixture in game_week.fixtures.all():
        mp = preds.get(fixture.id)
        if mp is None or mp.pred_home is None or mp.pred_away is None:
            unscored += 1
    if unscored:
        missing.append(f"{unscored} match score(s)")

    tgp = getattr(entry, "total_goals_prediction", None)
    if tgp is None or tgp.predicted_total is None:
        missing.append("total goals")

    named = sum(1 for p in entry.scorer_picks.all() if (p.player_name or "").strip())
    if named < 4:
        missing.append(f"{4 - named} scorer pick(s)")

    return entry, missing


def _shade(cell, hex_fill):
    """Apply a background fill to a table cell (python-docx has no direct API)."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_fill)
    cell._tc.get_or_add_tcPr().append(shd)


def _bold(cell):
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.bold = True


def _first_word(name: str) -> str:
    """First word of a team name ('Hull City' -> 'Hull')."""
    return (name or "").split()[0] if (name or "").strip() else (name or "")


# Clubs whose recognisable short form isn't the first word — nicknames (Aston
# Villa -> Villa) or shared-prefix clubs that must keep enough to tell them apart
# (Man City / Man Utd). Keyed by the full name, lowercased. Extend as needed.
_TEAM_SHORT = {
    "aston villa": "Villa",
    "crystal palace": "Palace",
    "nottingham forest": "Forest",
    "notts forest": "Forest",
    "manchester city": "Man City",
    "man city": "Man City",
    "manchester united": "Man Utd",
    "man united": "Man Utd",
    "man utd": "Man Utd",
    "west ham": "West Ham",
    "west ham united": "West Ham",
    "west bromwich albion": "West Brom",
    "west brom": "West Brom",
    "sheffield united": "Sheff Utd",
    "sheffield utd": "Sheff Utd",
    "sheffield wednesday": "Sheff Wed",
}


def _short_candidate(name: str) -> str:
    """Best short form for one team: explicit override, else first word."""
    key = " ".join((name or "").split()).lower()
    override = _TEAM_SHORT.get(key)
    if override:
        # Match the source's case so it sits consistently with the other teams.
        return override.upper() if (name or "").isupper() else override
    return _first_word(name)


def _team_shortener(names):
    """Return a function that shortens each team name, but falls back to the full
    name when two distinct teams that week would produce the same short form (so
    they never become indistinguishable on the sheet)."""
    from collections import defaultdict

    by_short = defaultdict(set)
    for n in names:
        by_short[_short_candidate(n).lower()].add((n or "").strip())

    def short(name):
        if not (name or "").strip():
            return name or ""
        candidate = _short_candidate(name)
        if len(by_short[candidate.lower()]) > 1:
            return name  # collision — keep the full name
        return candidate

    return short


def _scorer_label(pick) -> str:
    """Surname + team for a scorer pick, e.g. 'Haaland (Man City)'.

    Prefers the resolved Player (surname of full_name + club); falls back to
    parsing the raw typed text of the form 'First Last (Team)'.
    """
    raw = (pick.player_name or "").strip()
    if not raw and pick.player is None:
        return ""

    name, team = raw, ""
    if raw.endswith(")") and "(" in raw:
        name, _, tail = raw.partition("(")
        name, team = name.strip(), tail[:-1].strip()

    player = pick.player
    if player is not None:
        name = (player.full_name or name).strip()
        team = (player.club or team).strip()

    surname = name.split()[-1] if name.split() else name
    return f"{surname} ({team})" if team else surname


def _fixed_col_widths(table, widths_cm):
    """Pin per-column widths (cm) with a fixed table layout so Word honours them."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm

    table.autofit = False
    table.allow_autofit = False
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    table._tbl.tblPr.append(layout)
    for row in table.rows:
        for idx, w in enumerate(widths_cm):
            row.cells[idx].width = Cm(w)


def _row_heights(table, cm, skip_header=True):
    """Give rows a minimum height (cm) so the score boxes are easy to mark by
    hand. AT_LEAST means the row still grows if the text needs more room."""
    from docx.enum.table import WD_ROW_HEIGHT_RULE
    from docx.shared import Cm

    rows = table.rows[1:] if skip_header else table.rows
    for r in rows:
        r.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        r.height = Cm(cm)


def _legend(doc, text):
    """A small-print legend line with tight spacing."""
    from docx.shared import Pt

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.size = Pt(8)
    return p


def _heading(doc, text):
    """A compact bold section heading."""
    from docx.shared import Pt

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    p.add_run(text).bold = True
    return p


def build_entry_docx(participant, game_week) -> bytes:
    """Render `participant`'s entry for `game_week` as a .docx byte string."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Pt

    entry = (
        Entry.objects.filter(participant=participant, game_week=game_week)
        .prefetch_related(
            "match_predictions", "tf_answers", "scorer_picks"
        )
        .first()
    )
    preds = {mp.fixture_id: mp for mp in entry.match_predictions.all()} if entry else {}
    tf = {a.question_id: a.answer for a in entry.tf_answers.all()} if entry else {}
    scorers = {p.position: p for p in entry.scorer_picks.all()} if entry else {}
    tgp = getattr(entry, "total_goals_prediction", None) if entry else None

    doc = Document()

    # Compact everything so the whole form lands on a single A4 page: narrow
    # margins, smaller base font, no inter-paragraph spacing.
    normal = doc.styles["Normal"]
    normal.font.size = Pt(10)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.0
    for section in doc.sections:
        section.top_margin = Cm(1.0)
        section.bottom_margin = Cm(1.0)
        section.left_margin = Cm(1.3)
        section.right_margin = Cm(1.3)

    # --- Header: "Week N" + team name ----------------------------------------
    head = doc.add_paragraph()
    head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    head.paragraph_format.space_after = Pt(0)
    r = head.add_run(f"Week {game_week.week_number}")
    r.bold = True
    r.font.size = Pt(16)
    team = doc.add_paragraph()
    team.alignment = WD_ALIGN_PARAGRAPH.CENTER
    team.paragraph_format.space_after = Pt(6)
    r = team.add_run(participant.display_name)
    r.bold = True
    r.font.size = Pt(13)

    # --- Section 1: match scores (Home | H | Away | A | Score) ----------------
    fixtures = list(game_week.fixtures.all())
    short = _team_shortener(
        [f.home_team for f in fixtures] + [f.away_team for f in fixtures]
    )
    t1 = doc.add_table(rows=1, cols=5)
    t1.style = "Table Grid"
    for cell, text in zip(t1.rows[0].cells, ["Home", "", "Away", "", "Score"]):
        cell.text = text
        _bold(cell)
        _shade(cell, "D9D9D9")
    for fixture in fixtures:
        mp = preds.get(fixture.id)
        h = "" if not mp or mp.pred_home is None else str(mp.pred_home)
        a = "" if not mp or mp.pred_away is None else str(mp.pred_away)
        cells = t1.add_row().cells
        cells[0].text = short(fixture.home_team)
        cells[1].text = h
        cells[2].text = short(fixture.away_team)
        cells[3].text = a
        cells[4].text = ""  # scoring box, left blank for marking
    _fixed_col_widths(t1, [5.0, 1.1, 5.0, 1.1, 1.6])
    _row_heights(t1, 0.65)
    _legend(doc, _SECTION1_LEGEND)

    # --- Section 2: total goals ----------------------------------------------
    _heading(doc, "Total Number of Goals Scored in these 10 matches")
    t2 = doc.add_table(rows=1, cols=2)
    t2.style = "Table Grid"
    t2.rows[0].cells[0].text = (
        "" if not tgp or tgp.predicted_total is None else str(tgp.predicted_total)
    )
    t2.rows[0].cells[1].text = ""  # score box
    _fixed_col_widths(t2, [1.6, 1.6])
    _row_heights(t2, 0.65, skip_header=False)
    _legend(doc, _SECTION2_LEGEND)

    # --- Section 3: true / false ---------------------------------------------
    _heading(doc, "True / False (20 points for all 8 correct)")
    questions = list(game_week.questions.all())
    t3 = doc.add_table(rows=1, cols=4)
    t3.style = "Table Grid"
    for cell, text in zip(t3.rows[0].cells, ["Question", "True", "False", "Score"]):
        cell.text = text
        _bold(cell)
        _shade(cell, "D9D9D9")
    for q in questions:
        ans = tf.get(q.id)
        cells = t3.add_row().cells
        cells[0].text = q.text
        cells[1].text = "X" if ans is True else ""
        cells[2].text = "X" if ans is False else ""
        cells[3].text = ""
    _fixed_col_widths(t3, [10.8, 1.3, 1.3, 1.6])
    _row_heights(t3, 0.65)
    _legend(
        doc,
        "2pts for each correct answer and a bonus of 4pts if you get them all correct",
    )

    # --- Section 4: scorers ---------------------------------------------------
    p4 = _heading(doc, "")
    p4.add_run("Predict the Scorers: ").bold = True
    p4.add_run("bonus of 1pt for any extra goals scored by your selection").font.size = Pt(8)
    t4 = doc.add_table(rows=1, cols=3)
    t4.style = "Table Grid"
    for cell, text in zip(t4.rows[0].cells, ["Points", "Scorer and Team", "Score"]):
        cell.text = text
        _bold(cell)
        _shade(cell, "D9D9D9")
    for position in range(1, 5):
        cells = t4.add_row().cells
        pick = scorers.get(position)
        cells[0].text = _SECTION4_POINTS[position]
        cells[1].text = _scorer_label(pick) if pick else ""
        cells[2].text = ""
    _fixed_col_widths(t4, [2.0, 10.9, 1.6])
    _row_heights(t4, 0.65)

    # --- Grand total (blank box for marking) ---------------------------------
    _heading(doc, "")
    gt = doc.add_table(rows=1, cols=2)
    gt.style = "Table Grid"
    gt.rows[0].cells[0].text = "GRAND TOTAL"
    _bold(gt.rows[0].cells[0])
    _shade(gt.rows[0].cells[0], "D9D9D9")
    gt.rows[0].cells[1].text = ""
    _fixed_col_widths(gt, [3.6, 1.6])
    _row_heights(gt, 0.65, skip_header=False)

    return _to_bytes(doc)


def _to_bytes(doc) -> bytes:
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
