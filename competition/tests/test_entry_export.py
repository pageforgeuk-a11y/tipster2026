"""Download / email 'my entry' as a Word (.docx) form."""

import io
from datetime import timedelta

from django.contrib.auth.models import User
from django.core import mail
from django.test import TestCase, override_settings
from django.urls import reverse
from django.utils import timezone

from competition import entry_doc
from competition.models import (
    Entry,
    Fixture,
    GameWeek,
    MatchPrediction,
    Participant,
    ScorerPick,
    Season,
    TotalGoalsPrediction,
    TrueFalseAnswer,
    TrueFalseQuestion,
)


class EntryExportTests(TestCase):
    def setUp(self):
        self.season = Season.objects.create(name="S", is_active=True)
        self.user = User.objects.create_user("me@example.com", "me@example.com", "pw")
        self.participant = Participant.objects.create(
            user=self.user, season=self.season, display_name="Red Lion Rovers", join_week=1
        )
        self.gw = GameWeek.objects.create(
            season=self.season,
            week_number=1,
            title="THE CHARITY TIPSTER",
            date_range_label="Fri 21 Aug – Mon 24 Aug 2026",
            deadline=timezone.now() + timedelta(days=1),
            status=GameWeek.Status.OPEN,
        )
        self.fixtures = [
            Fixture.objects.create(
                game_week=self.gw, order=i, home_team=f"Home {i}", away_team=f"Away {i}"
            )
            for i in range(1, 11)
        ]
        self.questions = [
            TrueFalseQuestion.objects.create(game_week=self.gw, order=i, text=f"Q{i}?")
            for i in range(1, 9)
        ]
        self.client.force_login(self.user)

    def _fill_entry(self, *, scorers=4):
        entry = Entry.objects.create(
            participant=self.participant, game_week=self.gw, submitted_at=timezone.now()
        )
        for i, fx in enumerate(self.fixtures):
            MatchPrediction.objects.create(entry=entry, fixture=fx, pred_home=i % 4, pred_away=1)
        TotalGoalsPrediction.objects.create(entry=entry, predicted_total=25)
        for q in self.questions:
            TrueFalseAnswer.objects.create(entry=entry, question=q, answer=(q.order % 2 == 0))
        for pos in range(1, scorers + 1):
            ScorerPick.objects.create(entry=entry, position=pos, player_name=f"Scorer {pos}")
        return entry

    def test_completeness_reports_missing(self):
        _, missing = entry_doc.entry_completeness(self.participant, self.gw)
        self.assertEqual(missing, ["Save your predictions first"])
        self._fill_entry(scorers=2)
        _, missing = entry_doc.entry_completeness(self.participant, self.gw)
        self.assertEqual(missing, ["2 scorer pick(s)"])

    def test_download_returns_docx(self):
        self._fill_entry()
        resp = self.client.get(
            reverse("download_entry", args=[self.gw.week_number]), SERVER_NAME="localhost"
        )
        self.assertEqual(resp.status_code, 200)
        self.assertIn("wordprocessingml", resp["Content-Type"])
        self.assertIn("attachment", resp["Content-Disposition"])
        self.assertIn('"Tipsters WK1 - Red Lion Rovers.docx"', resp["Content-Disposition"])
        self.assertTrue(resp.content.startswith(b"PK"))  # .docx is a zip

    def test_entry_filename_format_and_sanitising(self):
        self.assertEqual(
            entry_doc.entry_filename(self.participant, self.gw),
            "Tipsters WK1 - Red Lion Rovers.docx",
        )
        self.participant.display_name = 'A/B:C team'  # illegal filename chars
        self.assertEqual(
            entry_doc.entry_filename(self.participant, self.gw),
            "Tipsters WK1 - ABC team.docx",
        )

    def test_download_blocked_when_incomplete(self):
        self._fill_entry(scorers=1)  # not all four scorers
        resp = self.client.get(
            reverse("download_entry", args=[self.gw.week_number]), SERVER_NAME="localhost"
        )
        self.assertRedirects(resp, reverse("entry", args=[self.gw.week_number]))

    @override_settings(
        EMAIL_PROVIDER="console",
        ORGANISER_EMAIL="neil@example.org",
        DEFAULT_FROM_EMAIL="Tipsters <tipsters@mail.pageforge.co.uk>",
    )
    def test_email_sends_to_organiser_with_attachment(self):
        self.user.first_name, self.user.last_name = "John", "Smith"
        self.user.save()
        self._fill_entry()
        resp = self.client.post(
            reverse("email_entry", args=[self.gw.week_number]), SERVER_NAME="localhost"
        )
        self.assertRedirects(resp, reverse("entry", args=[self.gw.week_number]))
        self.assertEqual(len(mail.outbox), 1)
        msg = mail.outbox[0]
        self.assertEqual(msg.to, ["neil@example.org"])  # organiser, not the player
        self.assertEqual(msg.subject, "WK1 Entry Red Lion Rovers")
        self.assertIn("Hi Neil", msg.body)
        self.assertIn("Red Lion Rovers Week 1", msg.body)
        # From shows the player's name, address stays on the verified domain.
        self.assertEqual(msg.from_email, "John Smith <tipsters@mail.pageforge.co.uk>")
        self.assertEqual(msg.reply_to, ["me@example.com"])  # reply goes to the player
        self.assertEqual(msg.attachments[0][0], "Tipsters WK1 - Red Lion Rovers.docx")

    def test_docx_contains_predictions(self):
        self._fill_entry()
        # Realistic values to exercise team-name shortening + scorer formatting.
        fx = self.fixtures[0]
        fx.home_team, fx.away_team = "Hull City", "Manchester United"
        fx.save()
        pick = ScorerPick.objects.get(entry__participant=self.participant, position=1)
        pick.player_name = "Erling Haaland (Man City)"
        pick.save()

        from docx import Document

        data = entry_doc.build_entry_docx(self.participant, self.gw)
        doc = Document(io.BytesIO(data))
        text = "\n".join(p.text for p in doc.paragraphs)
        cells = " ".join(c.text for t in doc.tables for row in t.rows for c in row.cells)
        self.assertIn("Week 1", text)  # header
        self.assertIn("Red Lion Rovers", text)  # team name in header
        self.assertIn("Hull", cells)  # shortened to first word
        self.assertNotIn("Hull City", cells)
        self.assertIn("Haaland (Man City)", cells)  # surname + team
        self.assertIn("25", cells)  # total goals

    def test_team_shortener(self):
        short = entry_doc._team_shortener(
            ["Hull City", "Everton", "Man City", "Man Utd", "Aston Villa", "Crystal Palace"]
        )
        self.assertEqual(short("Hull City"), "Hull")  # first word
        self.assertEqual(short("Everton"), "Everton")
        self.assertEqual(short("Aston Villa"), "Villa")  # nickname override
        self.assertEqual(short("Crystal Palace"), "Palace")
        self.assertEqual(short("Man City"), "Man City")  # kept distinct
        self.assertEqual(short("Man Utd"), "Man Utd")

    def test_team_shortener_collision_falls_back_to_full(self):
        # Two clubs that share a first word and have no override stay full.
        short = entry_doc._team_shortener(["Boston Rangers", "Boston Rovers"])
        self.assertEqual(short("Boston Rangers"), "Boston Rangers")
        self.assertEqual(short("Boston Rovers"), "Boston Rovers")

    def test_duplicate_scorer_text_rejected(self):
        from competition.forms import EntryForm

        data = {"scorer_1": "Erling Haaland (Man City)", "scorer_2": "Erling Haaland (Man City)"}
        form = EntryForm(data, fixtures=self.fixtures, questions=self.questions)
        self.assertFalse(form.is_valid())
        self.assertIn("scorer_2", form.errors)

    def test_duplicate_scorer_by_resolved_identity_rejected(self):
        from competition.forms import EntryForm
        from competition.models import Player

        Player.objects.create(full_name="Mohamed Salah", club="Liverpool")
        # Same player expressed two ways -> resolves to one identity.
        data = {"scorer_1": "Mohamed Salah (Liverpool)", "scorer_2": "Salah"}
        form = EntryForm(data, fixtures=self.fixtures, questions=self.questions)
        self.assertFalse(form.is_valid())
        self.assertIn("scorer_2", form.errors)

    def test_distinct_scorers_allowed(self):
        from competition.forms import EntryForm

        data = {"scorer_1": "Player A", "scorer_2": "Player B"}
        form = EntryForm(data, fixtures=self.fixtures, questions=self.questions)
        self.assertTrue(form.is_valid(), form.errors)

    def test_scorer_label_uses_resolved_player(self):
        from competition.models import Player

        player = Player.objects.create(full_name="Bukayo Saka", club="Arsenal")
        entry = self._fill_entry()
        pick = entry.scorer_picks.get(position=1)
        pick.player = player
        pick.player_name = "saka"  # raw typed text, no team
        pick.save()
        self.assertEqual(entry_doc._scorer_label(pick), "Saka (Arsenal)")
